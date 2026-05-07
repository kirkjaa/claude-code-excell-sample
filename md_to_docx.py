"""Convert the LDC Client Version proposal markdown to a formatted .docx file."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path(r"G:\Claude Cowork Retail\LDC Consulting Proposal - Client Version.md")
DST = Path(r"G:\Claude Cowork Retail\LDC Consulting Proposal - Client Version.docx")

INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_inline(paragraph, text):
    """Add text with inline **bold**, *italic*, `code` formatting."""
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for level, size in [(1, 20), (2, 16), (3, 13)]:
        h = styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def add_table_from_md(doc, header, rows):
    cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, text in enumerate(header):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        add_inline(p, text)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[i], "1F3A5F")
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for i, text in enumerate(row):
            if i >= cols:
                break
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            add_inline(p, text)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP


def parse_table(lines, idx):
    """Parse a markdown table starting at lines[idx]. Returns (header, rows, next_idx)."""
    def split_row(line):
        line = line.strip()
        if line.startswith("|"):
            line = line[1:]
        if line.endswith("|"):
            line = line[:-1]
        return [c.strip() for c in line.split("|")]

    header = split_row(lines[idx])
    # idx+1 is the separator line
    rows = []
    j = idx + 2
    while j < len(lines) and lines[j].lstrip().startswith("|"):
        row = split_row(lines[j])
        rows.append(row)
        j += 1
    return header, rows, j


def convert(src_path, dst_path):
    text = src_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    configure_styles(doc)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            p = doc.add_paragraph()
            p_pr = p._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "999999")
            p_bdr.append(bottom)
            p_pr.append(p_bdr)
            i += 1
            continue

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        # Tables
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|[\s\-:|]+\|\s*$", lines[i + 1]):
            header, rows, next_idx = parse_table(lines, i)
            add_table_from_md(doc, header, rows)
            doc.add_paragraph()
            i = next_idx
            continue

        # Bullet list
        if re.match(r"^\s*[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            content = re.sub(r"^\s*[-*]\s+", "", stripped)
            add_inline(p, content)
            i += 1
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            content = re.sub(r"^\s*\d+\.\s+", "", stripped)
            add_inline(p, content)
            i += 1
            continue

        # Default paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(str(dst_path))
    print(f"Wrote: {dst_path}")
    print(f"Size: {dst_path.stat().st_size:,} bytes")


if __name__ == "__main__":
    convert(SRC, DST)
